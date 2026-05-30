"""
Build the presentation script as a .docx file.
Three sections: Demo / Code / Tradeoffs.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/wang/Desktop/final-project-ada/HISTORY_LOG_presentation_script.docx"

doc = Document()

# ------------------------------------------------------------
# Base styles
# ------------------------------------------------------------
styles = doc.styles

# Normal body
normal = styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)


# ------------------------------------------------------------
# Helpers
# ------------------------------------------------------------
def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)


def add_subtitle(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def add_section_heading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.bold = True
    r.font.color.rgb = RGBColor(0x8b, 0x4a, 0x1c)


def add_subheading(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.bold = True


def add_timestamp(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = RGBColor(0x8b, 0x4a, 0x1c)


def add_cue(text):
    """A 'CLICK' / stage-direction line — bordered look via inline style."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("► " + text)
    r.font.name = 'Consolas'
    r.font.size = Pt(10)
    r.bold = True
    r.font.color.rgb = RGBColor(0x2a, 0x6a, 0x9a)


def add_say(text):
    """Spoken text — body paragraphs."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.35
    r = p.add_run(text)
    r.font.size = Pt(11)


def add_note(text):
    """An inline note / tip in italics."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6a, 0x6a, 0x6a)


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        for chunk in item:
            r = p.add_run(chunk[0])
            r.font.size = Pt(11)
            if 'bold' in chunk[1]: r.bold = True
            if 'italic' in chunk[1]: r.italic = True
            if 'mono' in chunk[1]:
                r.font.name = 'Consolas'
                r.font.size = Pt(10)


def add_rule():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('— — — — — — — — — — — — — — — — — — — — — — — — — — — — — — —')
    r.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)
    r.font.size = Pt(9)


# ============================================================
# TITLE BLOCK
# ============================================================
add_title("HISTORY.LOG — Presentation Script")
add_subtitle("A 10-minute walkthrough of the final project: demo · code · tradeoffs")

p = doc.add_paragraph()
r = p.add_run("Zeyu Wang  ·  May 2026")
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_rule()

# Pre-flight
add_subheading("Pre-flight (before you hit record)")
add_bullets([
    [("Open the site on the cover page. Refresh once so MathJax is loaded.", "")],
    [("Have the left spine (TOC) visible — that's your jump-anywhere navigation.", "")],
    [("Quick test: flip one page forward and back to confirm the scanline wipe is smooth.", "")],
    [("Keep this script open on a second screen, or printed.", "")],
])

# Timing budget table
add_subheading("Budget at a glance")
add_bullets([
    [("Section 1 — Demo: ~5:00", "bold")],
    [("Section 2 — Code walkthrough: ~2:30", "bold")],
    [("Section 3 — Tradeoffs: ~2:30", "bold")],
])

doc.add_page_break()


# ============================================================
# SECTION 1 — DEMO
# ============================================================
add_section_heading("Section 1 · Demo  (0:00 – 5:00)")

# --- 0:00–0:45 — Open ---
add_timestamp("[0:00 – 0:45]  Opening — what is this thing?")
add_cue("You are on the cover page.")
add_say(
    "So I built this to walk through how AI got to where it is today. "
    "The question I kept asking myself while reading about all these models was — "
    "how do you actually go from typing rules into a computer to ChatGPT? "
    "Like, what's the actual chain of ideas?"
)
add_say(
    "It turns out the story is pretty clean. It's basically the same question "
    "being answered differently for sixty years. Every chapter on this timeline is one attempt at it."
)
add_say(
    "The look is meant to feel like an old computer manual — the kind you'd find "
    "next to a DEC terminal. There are fifteen chapters; I'll walk you through five."
)

# --- 0:45–1:30 — Symbolic ---
add_timestamp("[0:45 – 1:30]  Chapter 01 — Symbolic AI")
add_cue("Click '01 1956 SYMBOLIC AI' in the spine.")
add_say(
    "That's the page-flip — every chapter is its own page you can flip through."
)
add_say(
    "So this is where AI starts. 1956. The idea is: if I write down enough if-then rules, "
    "I can make a computer think. The famous example is MYCIN, a medical system from the '70s — "
    "about 600 rules for diagnosing infections. It actually worked pretty well."
)
add_cue("Scroll to 'Try it yourself.' Make sure the '20 questions' tab is selected.")
add_say(
    "Here's a tiny version. The system asks me questions to identify the animal I'm thinking of. "
    "Let me think of a dog."
)
add_cue("Click YES on 'Does it have fur?'")
add_say(
    "Each question is picked to split the candidates as evenly as possible — that's the H number, "
    "the information gain. So it narrows down fast."
)
add_cue("Answer two or three more questions until it lands on Dog.")
add_say(
    "Now the catch. If I think of a kangaroo, which isn't in the database, this system "
    "has nothing to say. That's the limit that killed this approach — the real world has "
    "too many exceptions to write down by hand."
)

# --- 1:30–3:30 — Perceptron ---
add_timestamp("[1:30 – 3:30]  Chapter 02 — Perceptron + XOR proof")
add_cue("Press → or click chapter 02 in the spine.")
add_say(
    "The alternative is: don't write the rules. Learn them from data. "
    "That's the Perceptron, 1957."
)
add_cue("Point at the equation in section 2.")
add_say(
    "The math is one line. Take a weighted sum of the inputs, check if it's positive or negative — "
    "that's your prediction. And you learn the weights by nudging them whenever you're wrong."
)
add_cue("Scroll to 'Try it yourself,' click the 'linearly separable' preset.")
add_say(
    "Let me show you. Two clusters — orange class and blue class."
)
add_cue("Press '▶ Train (animated).'")
add_say(
    "Watch the orange line. Every step, the perceptron looks at the misclassified points — "
    "the ones circled in red — and pushes the line to fix them. The green dashed arrow is "
    "the weight vector; it rotates as the model learns. Mistakes drop to zero. Converged."
)
add_say(
    "This was huge in 1957. People thought we'd basically figured out intelligence. "
    "The New York Times reported the Navy expected this machine to walk, talk, and reproduce itself."
)
add_cue("Now press the 'XOR (impossible)' preset.")
add_say(
    "Then 1969 happened. Minsky and Papert proved this can't learn this pattern — XOR. "
    "Four points, alternating corners."
)
add_cue("Press '▶ Train.'")
add_say(
    "Watch — it never converges. The line just oscillates. The mistake count never reaches zero."
)
add_cue("Scroll up to the 'Why XOR is impossible' equation block, click to expand the proof.")
add_say(
    "The proof is short. You assume there's a line that works, write down the four inequalities — "
    "one per point — and they contradict each other. So no such line exists."
)
add_say(
    "This single result killed neural networks for fifteen years. AI funding dried up. "
    "People call it the first AI winter."
)

# --- 3:30–4:30 — MLP ---
add_timestamp("[3:30 – 4:30]  Chapter 03 — MLP solves XOR")
add_cue("Press → or click chapter 03.")
add_say(
    "The fix existed in theory all along — you just stack perceptrons. Layer on top of another, "
    "with a non-linear function in between. The blocker was that nobody knew how to train "
    "the middle layer."
)
add_say(
    "The answer is backpropagation — literally just the chain rule from calculus, applied carefully. "
    "Errors at the output flow backward through the weights, every layer gets a usable gradient. "
    "The paper that popularised this in 1986 was by Hinton — the same Hinton who just won the Nobel."
)
add_cue("Scroll to 'Try it yourself,' press '▶ Train.'")
add_say(
    "The shaded background is the network's prediction across the plane — orange where it predicts 1, "
    "blue where it predicts 0. As gradient descent runs, the boundary bends and wraps around the points. "
    "Within a few seconds it solves XOR — the same problem that broke the perceptron — because the "
    "hidden layer learned to rewrite the input into a new representation where one line works."
)
add_note("If time is tight, skip the H = 1 demo and move on.")

# --- 4:30–5:00 — Attention ---
add_timestamp("[4:30 – 5:00]  Chapter 08 — Attention (the modern leap)")
add_cue("Click chapter 08 'ATTENTION' in the spine.")
add_say(
    "I'm going to skip a few chapters and jump to 2014. Attention."
)
add_say(
    "The problem it solves: the old approach squashed an entire input sentence into one fixed vector. "
    "For long sentences, you lose a lot. Attention says: don't squash. At each step, look back at "
    "every word and weigh them by relevance."
)
add_cue("Scroll to 'Try it yourself' (default sentence already selected).")
add_say(
    "Classic example. The word 'it' — what does it refer to?"
)
add_cue("Click the token 'it.'")
add_say(
    "Most of the attention weight is on 'animal.' The model has learned to look back to the right word. "
    "This idea, generalised, became the Transformer in 2017 — the foundation of every modern LLM."
)

add_rule()


# ============================================================
# SECTION 2 — CODE
# ============================================================
add_section_heading("Section 2 · Code  (5:00 – 7:30)")

# --- 5:00–5:30 — overview ---
add_timestamp("[5:00 – 5:30]  Architecture overview")
add_cue("Open your editor side-by-side, show the project folder.")
add_say(
    "The whole site is five files. No build step, no framework, no backend. "
    "You open index.html and it runs."
)
add_bullets([
    [("index.html  ", "mono"), ("— the shell. Three regions: spine (TOC), pages, footer.", "")],
    [("styles.css  ", "mono"), ("— the vintage terminal aesthetic. Amber phosphor, scanlines, the page-flip wipe.", "")],
    [("data.js  ", "mono"), ("— all the content. Each chapter is one big object: title, narrative, math, code snippet, papers.", "")],
    [("app.js  ", "mono"), ("— the renderer and the page-flip state machine.", "")],
    [("interactives.js  ", "mono"), ("— fifteen demo functions. Each one is mounted lazily when you open its chapter.", "")],
])

# --- 5:30–6:15 — data shape ---
add_timestamp("[5:30 – 6:15]  How content is structured — data.js")
add_cue("Show data.js. Scroll to the 'perceptron' entry.")
add_say(
    "Each chapter is one object that fully describes itself. Title, year, key idea, a narrative paragraph, "
    "the equations — and crucially, every equation has an optional 'proof' field with the step-by-step "
    "derivation. That's what powers the click-to-reveal proof. The math is data, not hand-built HTML."
)
add_say(
    "Each chapter also lists the original papers — author, venue, year, URL, and an optional local file path "
    "pointing into the 'papers/' folder. That's how the inline [1] citations link to the bottom of the chapter."
)

# --- 6:15–6:55 — app.js ---
add_timestamp("[6:15 – 6:55]  app.js — page-flip + math + TOC")
add_cue("Open app.js.")
add_say(
    "Three responsibilities. First, render every chapter as a hidden div. Second, maintain a 'Book' "
    "state machine — current page, total pages — and animate the flip when you press next or previous. "
    "The flip itself is just two CSS keyframes: clip-path slides off the top, the next page slides on "
    "from the bottom. Pure CSS, no library."
)
add_say(
    "Third, the math click-for-proof. When you click an equation block, it toggles a class, the hidden "
    "proof div unfolds, and MathJax re-typesets just that block."
)
add_say(
    "Interactives are mounted lazily — the moment a page is about to be shown, app.js calls "
    "window.Demos[kind] to instantiate that demo. So the cost of opening the page is one chapter at a time."
)

# --- 6:55–7:30 — interactives.js ---
add_timestamp("[6:55 – 7:30]  interactives.js — fifteen self-contained demos")
add_cue("Open interactives.js, scroll to Demos.perceptron.")
add_say(
    "Each demo is a single function that takes a host element and fills it. The perceptron demo, for example, "
    "is a few hundred lines: a canvas, click handlers, an animated training loop driven by setInterval, "
    "live status updates. The MLP demo runs actual backprop on XOR — the orange-and-blue contour you see "
    "is the network's live output evaluated on a grid of points after every batch."
)
add_say(
    "Everything is vanilla — Canvas 2D for plots, no D3, no Chart.js. Keeps the bundle at zero bytes."
)

add_rule()


# ============================================================
# SECTION 3 — TRADEOFFS
# ============================================================
add_section_heading("Section 3 · Tradeoffs  (7:30 – 10:00)")

# --- 7:30–8:00 — framework choice ---
add_timestamp("[7:30 – 8:00]  Vanilla JS, no framework")
add_say(
    "First decision: React, Next.js, or plain HTML. I went with plain HTML."
)
add_say(
    "The reason is that the demos don't share state — each one is its own little world. "
    "A framework would add a build step, a deploy story, and a hundred dependencies, "
    "and I'd gain almost nothing from the component model. The cost was that I had to write "
    "DOM manipulation by hand, which is verbose. But the file you open in the browser is the "
    "file I wrote — no transpilation, no minification, no surprises."
)

# --- 8:00–8:30 — client-only ---
add_timestamp("[8:00 – 8:30]  Client-only — no backend")
add_say(
    "Second decision: should the GPT and BERT demos call a real API? I chose no."
)
add_say(
    "If I called OpenAI from the page, I'd need a server to hold the key, which means hosting, "
    "rate limits, and a cost per page-load. So the language-model demos use precomputed examples — "
    "the masked-LM predictions for BERT, the side-by-side base-vs-aligned outputs for RLHF. "
    "You lose the magic of typing your own prompt, but you gain something more important — the site "
    "runs anywhere, forever, with no key, no internet for the demos to function."
)

# --- 8:30–9:00 — aesthetic ---
add_timestamp("[8:30 – 9:00]  Minimalist → vintage terminal")
add_say(
    "Third decision was visual. I built a clean minimalist version first — Garamond serif, "
    "lots of whitespace, very Stripe-docs. It worked, but it looked like every other tech-explainer site. "
    "I switched to the amber CRT aesthetic. It gives the project a personality. The tradeoff is that "
    "monospace everywhere is harder to read in long paragraphs, so I had to bump the line-height "
    "and tune the colors so the body text doesn't strain your eyes."
)

# --- 9:00–9:30 — page-flip ---
add_timestamp("[9:00 – 9:30]  Page-flip vs. infinite scroll")
add_say(
    "Fourth: how do you navigate fifteen chapters? My first version was one long scrollable page "
    "with collapsible entries — modern, easy to skim, but it discourages reading anything in full. "
    "I switched to one chapter per page with a flip animation. It feels more like a book; people "
    "actually finish a chapter before moving on. The cost is that it forces a sequence, "
    "so I added the spine on the left as a jump-anywhere index."
)

# --- 9:30–10:00 — fidelity ---
add_timestamp("[9:30 – 10:00]  Where to spend the simulation budget")
add_say(
    "Last one. For some demos — perceptron, MLP — the math is small enough that I run the real "
    "algorithm in the browser. Real backprop on XOR, ~2,000 epochs per second. For others — "
    "Transformer attention, BERT — running the real thing in JavaScript isn't practical, "
    "so I use hand-designed examples that make the underlying concept obvious."
)
add_say(
    "The principle I tried to stick to: simulate where it teaches something, precompute where it doesn't. "
    "A live perceptron failing on XOR teaches you something the diagram can't. A live tokeniser, on the "
    "other hand, mostly just teaches you that JavaScript is slow."
)

add_rule()

add_say("Thanks.")

# ============================================================
# Tips / contingencies on a final page
# ============================================================
doc.add_page_break()
add_section_heading("Tips & contingencies")

add_subheading("If a demo doesn't load")
add_say("Flip past it and narrate the static content. The cover page narration alone fills 30 seconds.")

add_subheading("If you're running long")
add_bullets([
    [("Drop the H = 1 demo in chapter 3 (the MLP one).", "")],
    [("Skip the second attention sentence in chapter 8.", "")],
    [("Combine tradeoffs 3 and 4 into one paragraph: 'visuals and navigation both got rewritten once.'", "")],
])

add_subheading("If you're running short (~30s slack)")
add_bullets([
    [("Click 'show proof' on Novikoff convergence in chapter 2 and read the four-line argument aloud.", "")],
    [("Add a beat after the XOR perceptron starts oscillating — let the audience watch it fail.", "")],
])

add_subheading("Things to lean into")
add_bullets([
    [("The 'kangaroo' line in Symbolic AI — concrete and casual.", "")],
    [("The pause when the XOR perceptron just keeps oscillating.", "")],
    [("The moment the MLP decision boundary starts curving — let it play for a beat.", "")],
])

add_subheading("Style reminders")
add_bullets([
    [("Don't read this verbatim. Skim it once, then talk through it from the click cues.", "")],
    [("Pauses are free. Let demos play.", "")],
    [("If you stumble, the demo is still doing the work — just look at the screen and narrate.", "")],
])

# ------------------------------------------------------------
doc.save(OUT)
print(f"wrote {OUT}")
